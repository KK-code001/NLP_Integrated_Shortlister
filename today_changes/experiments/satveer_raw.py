import docx

doc = docx.Document(r"c:\Users\Vivaan\Downloads\NLP_Internship\NLP_Integrated_Shortlister\ipdocs\Satveer_CV.docx")
lines = []
for p in doc.paragraphs:
    if p.text.strip():
        lines.append(p.text.strip())

for t in doc.tables:
    for r in t.rows:
        row_str = " | ".join(c.text.strip() for c in r.cells if c.text.strip())
        if row_str:
            lines.append(row_str)

with open(r"c:\Users\Vivaan\Downloads\NLP_Internship\NLP_Integrated_Shortlister\satveer_raw.txt", "w", encoding="utf-8") as f:
    f.write("\n".join(lines))

print("Wrote satveer_raw.txt successfully")
