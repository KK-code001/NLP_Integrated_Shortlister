import os
import re
import pandas as pd

class ScannedPDFError(Exception):
    """Raised when a PDF has no extractable text AND OCR also failed/was unavailable."""
    pass

class CorruptPDFError(Exception):
    """Raised when a file can't be opened/parsed at all."""
    pass

class UnsupportedFileTypeError(Exception):
    """Raised when the file extension isn't one we know how to read."""
    pass

def _extract_pdf_layout_aware(pdf_path: str) -> str:
    """
    Layout-aware text extraction using PyMuPDF (fitz).
    Sorts blocks spatially (top-to-bottom, left-to-right) so two-column
    resumes are read column-by-column rather than line-merged.
    Falls back to pdfplumber if PyMuPDF is unavailable.
    """
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(pdf_path)
        full_text = []
        for page in doc:
            blocks = page.get_text("blocks")
            # Sort blocks: y0 (top to bottom), x0 (left to right)
            blocks.sort(key=lambda b: (b[1], b[0]))
            for b in blocks:
                text = b[4].strip()
                if text:
                    full_text.append(text)
        return "\n".join(full_text)
    except Exception:
        # Fallback to pdfplumber
        import pdfplumber
        try:
            with pdfplumber.open(pdf_path) as pdf:
                text_parts = [page.extract_text() for page in pdf.pages]
            return "\n".join(t for t in text_parts if t)
        except Exception as e:
            raise CorruptPDFError(f"Could not open '{pdf_path}' as PDF ({e}).") from e

def _ocr_pdf(pdf_path: str) -> str:
    """OCR fallback for image-only/scanned PDFs using pdf2image and pytesseract."""
    try:
        from pdf2image import convert_from_path
        import pytesseract

        pages = convert_from_path(pdf_path)
        text_parts = [pytesseract.image_to_string(page) for page in pages]
        return "\n".join(t for t in text_parts if t)
    except Exception as e:
        return ""

def extract_text_from_pdf(pdf_path: str, allow_ocr: bool = True) -> str:
    full_text = _extract_pdf_layout_aware(pdf_path)

    if full_text.strip():
        return full_text

    if not allow_ocr:
        raise ScannedPDFError(f"No text layer found in '{pdf_path}' and OCR is disabled.")

    ocr_text = _ocr_pdf(pdf_path)
    if not ocr_text.strip():
        raise ScannedPDFError(f"No text found in '{pdf_path}', even after OCR.")
    return ocr_text

def extract_text_from_docx(docx_path: str) -> str:
    import docx
    try:
        doc = docx.Document(docx_path)
    except Exception as e:
        raise CorruptPDFError(f"Could not open '{docx_path}' as DOCX ({e}).") from e

    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)

    full_text = "\n".join(parts)
    if not full_text.strip():
        raise ScannedPDFError(f"No extractable text in '{docx_path}'.")
    return full_text

def extract_text_from_image(image_path: str) -> str:
    from PIL import Image
    import pytesseract

    try:
        img = Image.open(image_path)
    except Exception as e:
        raise CorruptPDFError(f"Could not open '{image_path}' as image ({e}).") from e

    text = pytesseract.image_to_string(img)
    if not text.strip():
        raise ScannedPDFError(f"No text found in image '{image_path}'.")
    return text

_EXTENSION_HANDLERS = {
    ".pdf": extract_text_from_pdf,
    ".docx": extract_text_from_docx,
    ".jpg": extract_text_from_image,
    ".jpeg": extract_text_from_image,
    ".png": extract_text_from_image,
    ".txt": lambda p: open(p, "r", encoding="utf-8", errors="ignore").read()
}

def extract_text_from_file(file_path: str) -> str:
    """
    Unified format-agnostic text extractor. Auto-detects extension (.pdf, .docx, .jpg, .png, .txt).
    """
    ext = os.path.splitext(file_path)[1].lower()
    handler = _EXTENSION_HANDLERS.get(ext)
    if handler is None:
        raise UnsupportedFileTypeError(
            f"'{ext}' is not supported. Supported: {', '.join(sorted(_EXTENSION_HANDLERS.keys()))}."
        )
    return handler(file_path)

def safe_extract_text_from_file(file_path: str):
    try:
        return extract_text_from_file(file_path), None
    except (ScannedPDFError, CorruptPDFError, UnsupportedFileTypeError) as e:
        return None, str(e)
