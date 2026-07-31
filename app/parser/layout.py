"""
Layout-aware document parser.

Priority order for PDF:
  1. Docling   — structured blocks, headings, tables, lists, reading order
  2. PyMuPDF   — two-column aware block sorting
  3. pdfplumber — plain text fallback
  4. OCR        — pytesseract for scanned/image-only PDFs

For DOCX:
  1. Docling
  2. python-docx (with paragraph style detection)

For images (.jpg, .png, etc.):
  1. pytesseract OCR

Every path returns a ResumeDocument with a blocks list and a raw_text string.
"""
from __future__ import annotations

import os
import logging
from dataclasses import dataclass, field
from typing import Optional

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Data model
# ---------------------------------------------------------------------------

@dataclass
class TextBlock:
    text: str
    block_type: str       # "heading" | "paragraph" | "list_item" | "table_cell" | "text"
    page: int = 0
    bbox: Optional[tuple] = None   # (x0, y0, x1, y1) — None for non-spatial sources
    is_bold: bool = False
    font_size: float = 0.0


@dataclass
class ResumeDocument:
    blocks: list[TextBlock] = field(default_factory=list)
    raw_text: str = ""            # Full concatenated text (compat. with feature_builder)
    source_path: str = ""
    parser_used: str = "unknown"
    parser_warnings: list[str] = field(default_factory=list)  # diagnostic log of parser chain


# ---------------------------------------------------------------------------
# 1. Docling parser
# ---------------------------------------------------------------------------

def _parse_with_docling(file_path: str) -> Optional[ResumeDocument]:
    """
    Use Docling to parse the document into a structured ResumeDocument.
    Handles PDF, DOCX, and other formats supported by Docling.
    Returns None if Docling is not installed or conversion fails.
    """
    try:
        from docling.document_converter import DocumentConverter
        logger.info("[Parser] Trying Docling for: %s", file_path)

        converter = DocumentConverter()
        result = converter.convert(file_path)
        doc = result.document

        blocks: list[TextBlock] = []

        # Iterate over all document items in reading order
        for element, _level in doc.iterate_items():
            # Resolve page number safely
            page_no = 0
            if hasattr(element, "prov") and element.prov:
                try:
                    page_no = element.prov[0].page_no
                except (AttributeError, IndexError):
                    pass

            element_text = ""
            if hasattr(element, "text"):
                element_text = (element.text or "").strip()

            if not element_text:
                continue

            # Classify by Docling element type (handle name changes across versions)
            cls_name = type(element).__name__

            if "SectionHeader" in cls_name or "Heading" in cls_name:
                blocks.append(TextBlock(
                    text=element_text,
                    block_type="heading",
                    page=page_no,
                ))

            elif "List" in cls_name:
                blocks.append(TextBlock(
                    text=element_text,
                    block_type="list_item",
                    page=page_no,
                ))

            elif "Table" in cls_name:
                # Try DataFrame export; fall back to markdown text
                try:
                    df = element.export_to_dataframe()
                    for _, row in df.iterrows():
                        for cell_val in row:
                            cell_text = str(cell_val).strip()
                            if cell_text and cell_text.lower() != "nan":
                                blocks.append(TextBlock(
                                    text=cell_text,
                                    block_type="table_cell",
                                    page=page_no,
                                ))
                except Exception:
                    # Fallback: get markdown of the table
                    try:
                        md = element.export_to_markdown()
                        if md:
                            blocks.append(TextBlock(
                                text=md.strip(),
                                block_type="table_cell",
                                page=page_no,
                            ))
                    except Exception:
                        pass

            elif "Text" in cls_name or "Paragraph" in cls_name:
                blocks.append(TextBlock(
                    text=element_text,
                    block_type="paragraph",
                    page=page_no,
                ))

            else:
                # Unknown element type — include as paragraph
                blocks.append(TextBlock(
                    text=element_text,
                    block_type="paragraph",
                    page=page_no,
                ))

        if not blocks:
            # Docling parsed something but yielded no text blocks — try markdown export
            try:
                md = doc.export_to_markdown()
                if md and md.strip():
                    return _parse_markdown_to_document(md, file_path, "docling_md")
            except Exception:
                pass
            return None

        raw_text = "\n".join(b.text for b in blocks if b.text)
        logger.info("[Parser] ✓ Docling succeeded — %d blocks extracted", len(blocks))
        return ResumeDocument(
            blocks=blocks,
            raw_text=raw_text,
            source_path=file_path,
            parser_used="docling",
        )

    except ImportError:
        logger.info("[Parser] ✗ Docling not installed — skipping")
        return None
    except Exception as exc:
        logger.warning("[Parser] ✗ Docling failed: %s — trying next parser", exc)
        return None


def _parse_markdown_to_document(
    md_text: str, source_path: str, parser_label: str
) -> ResumeDocument:
    """Convert markdown text into TextBlocks by interpreting # headings and bullet lists."""
    import re
    blocks: list[TextBlock] = []
    for line in md_text.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("#"):
            heading_text = stripped.lstrip("#").strip()
            blocks.append(TextBlock(text=heading_text, block_type="heading"))
        elif stripped.startswith(("-", "*", "•", "·")):
            item_text = re.sub(r"^[-*•·]\s*", "", stripped)
            blocks.append(TextBlock(text=item_text, block_type="list_item"))
        else:
            blocks.append(TextBlock(text=stripped, block_type="paragraph"))
    raw_text = "\n".join(b.text for b in blocks)
    return ResumeDocument(
        blocks=blocks, raw_text=raw_text,
        source_path=source_path, parser_used=parser_label,
    )


# ---------------------------------------------------------------------------
# 2. PyMuPDF fallback — two-column aware
# ---------------------------------------------------------------------------

def _parse_with_pymupdf(file_path: str) -> Optional[ResumeDocument]:
    """
    Extract text from PDF using PyMuPDF with two-column reconstruction.
    Blocks are sorted top-to-bottom in 20px row buckets, then left-to-right,
    so left-column text precedes right-column text on the same visual row.
    """
    try:
        import fitz  # PyMuPDF
        logger.info("[Parser] Trying PyMuPDF for: %s", file_path)

        doc_fitz = fitz.open(file_path)
        blocks: list[TextBlock] = []

        for page_num, page in enumerate(doc_fitz, start=1):
            raw_blocks = page.get_text("dict")["blocks"]

            # Bucket rows by y-position (20px tolerance) then sort x within bucket
            raw_blocks.sort(
                key=lambda b: (round(b["bbox"][1] / 20) * 20, b["bbox"][0])
            )

            for blk in raw_blocks:
                if blk.get("type") != 0:   # 0 = text, 1 = image
                    continue
                for line in blk.get("lines", []):
                    for span in line.get("spans", []):
                        span_text = (span.get("text") or "").strip()
                        if not span_text:
                            continue
                        flags = span.get("flags", 0)
                        is_bold = bool(flags & (1 << 4))
                        font_size = float(span.get("size", 0.0))
                        blocks.append(TextBlock(
                            text=span_text,
                            block_type="text",
                            page=page_num,
                            bbox=tuple(span.get("bbox", ())),
                            is_bold=is_bold,
                            font_size=font_size,
                        ))

        if not blocks:
            logger.info("[Parser] ✗ PyMuPDF found no text blocks (scanned PDF?)")
            return None

        raw_text = "\n".join(b.text for b in blocks if b.text)
        logger.info("[Parser] ✓ PyMuPDF succeeded — %d blocks extracted", len(blocks))
        return ResumeDocument(
            blocks=blocks,
            raw_text=raw_text,
            source_path=file_path,
            parser_used="pymupdf",
        )

    except ImportError:
        logger.info("[Parser] ✗ PyMuPDF (fitz) not installed — skipping")
        return None
    except Exception as exc:
        logger.warning("[Parser] ✗ PyMuPDF failed: %s — trying next parser", exc)
        return None


# ---------------------------------------------------------------------------
# 3. pdfplumber fallback
# ---------------------------------------------------------------------------

def _parse_with_pdfplumber(file_path: str) -> Optional[ResumeDocument]:
    try:
        import pdfplumber
        logger.info("[Parser] Trying pdfplumber for: %s", file_path)

        parts: list[str] = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    parts.append(t)

        if not parts:
            logger.info("[Parser] ✗ pdfplumber found no text")
            return None

        raw_text = "\n".join(parts)
        blocks = [
            TextBlock(text=line.strip(), block_type="paragraph")
            for line in raw_text.splitlines()
            if line.strip()
        ]
        logger.info("[Parser] ✓ pdfplumber succeeded — %d lines extracted", len(blocks))
        return ResumeDocument(
            blocks=blocks,
            raw_text=raw_text,
            source_path=file_path,
            parser_used="pdfplumber",
        )

    except Exception as exc:
        logger.warning("[Parser] ✗ pdfplumber failed: %s", exc)
        return None


# ---------------------------------------------------------------------------
# 4. python-docx parser
# ---------------------------------------------------------------------------

def _parse_docx(file_path: str) -> Optional[ResumeDocument]:
    try:
        import docx as python_docx

        document = python_docx.Document(file_path)
        blocks: list[TextBlock] = []

        for para in document.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style_name = (para.style.name or "").lower()
            if "heading" in style_name:
                block_type = "heading"
            elif "list" in style_name:
                block_type = "list_item"
            else:
                block_type = "paragraph"
            is_bold = any(run.bold for run in para.runs if run.bold is not None)
            blocks.append(TextBlock(
                text=text,
                block_type=block_type,
                is_bold=is_bold,
            ))

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        blocks.append(TextBlock(text=cell_text, block_type="table_cell"))

        if not blocks:
            return None

        raw_text = "\n".join(b.text for b in blocks)
        return ResumeDocument(
            blocks=blocks,
            raw_text=raw_text,
            source_path=file_path,
            parser_used="docx",
        )

    except Exception:
        return None


# ---------------------------------------------------------------------------
# 5. OCR fallback
# ---------------------------------------------------------------------------

def _parse_with_ocr(file_path: str) -> Optional[ResumeDocument]:
    try:
        logger.info("[Parser] Trying OCR (pytesseract) for: %s", file_path)
        ext = os.path.splitext(file_path)[1].lower()
        pages = []

        if ext == ".pdf":
            from pdf2image import convert_from_path
            pages = convert_from_path(file_path)
        else:
            from PIL import Image
            pages = [Image.open(file_path)]

        import pytesseract
        all_text = "\n".join(pytesseract.image_to_string(p) for p in pages)

        if not all_text.strip():
            logger.info("[Parser] ✗ OCR produced no text")
            return None

        blocks = [
            TextBlock(text=line.strip(), block_type="paragraph")
            for line in all_text.splitlines()
            if line.strip()
        ]
        logger.info("[Parser] ✓ OCR succeeded — %d lines extracted", len(blocks))
        return ResumeDocument(
            blocks=blocks,
            raw_text=all_text,
            source_path=file_path,
            parser_used="ocr",
        )

    except ImportError as exc:
        logger.warning("[Parser] ✗ OCR dependencies not installed (%s) — install pdf2image and pytesseract", exc)
        return None
    except Exception as exc:
        logger.warning("[Parser] ✗ OCR failed: %s", exc)
        return None


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def parse_resume_document(file_path: str) -> ResumeDocument:
    """
    Parse a resume file into a structured ResumeDocument.

    Format dispatch + fallback chain:
      PDF  → Docling → PyMuPDF → pdfplumber → OCR
      DOCX → Docling → python-docx
      Image (.jpg, .png, …) → OCR
      TXT  → plain read
    """
    ext = os.path.splitext(file_path)[1].lower()
    logger.info("[Parser] Parsing file: %s (format: %s)", file_path, ext)
    warnings: list[str] = []

    if ext in (".jpg", ".jpeg", ".png", ".bmp", ".tiff", ".webp"):
        doc = _parse_with_ocr(file_path)
        if doc:
            return doc
        warnings.append("OCR failed for image file")
        return _empty_document(file_path, "ocr_failed", warnings)

    if ext == ".txt":
        try:
            raw = open(file_path, "r", encoding="utf-8", errors="ignore").read()
            blocks = [
                TextBlock(text=line.strip(), block_type="paragraph")
                for line in raw.splitlines()
                if line.strip()
            ]
            logger.info("[Parser] ✓ TXT read — %d lines", len(blocks))
            return ResumeDocument(
                blocks=blocks, raw_text=raw,
                source_path=file_path, parser_used="txt",
            )
        except Exception as exc:
            warnings.append(f"TXT read failed: {exc}")
            return _empty_document(file_path, "txt_failed", warnings)

    if ext == ".docx":
        for parser_fn, name in [(_parse_docx, "python-docx"), (_parse_with_docling, "Docling")]:
            doc = parser_fn(file_path)
            if doc:
                return doc
            warnings.append(f"{name} failed for DOCX")
        return _empty_document(file_path, "docx_failed", warnings)

    if ext == ".pdf":
<<<<<<< HEAD
        return (
            _parse_with_docling(file_path)          # best: AI layout + reading order
            or _parse_with_pymupdf(file_path)       # good: column-aware geometric sort
            or _parse_with_pdfplumber(file_path)    # basic: plain text dump
            or _parse_with_ocr(file_path)           # last resort: scanned/image PDFs
            or _empty_document(file_path, "pdf_failed")
        )
=======
        for parser_fn, name in [
            (_parse_with_pymupdf, "PyMuPDF"),
            (_parse_with_pdfplumber, "pdfplumber"),
            (_parse_with_docling, "Docling"),
            (_parse_with_ocr, "OCR"),
        ]:
            doc = parser_fn(file_path)
            if doc:
                doc.parser_warnings = warnings  # attach chain history
                return doc
            warnings.append(f"{name} failed")

        logger.error("[Parser] ✗ ALL parsers failed for PDF: %s", file_path)
        return _empty_document(file_path, "pdf_failed", warnings)
>>>>>>> 5e3c81e71f02af9a7b0079f7930b2897de4273c9

    # Unsupported extension — try Docling anyway
    doc = _parse_with_docling(file_path)
    if doc:
        return doc
    warnings.append(f"Unsupported extension '{ext}' and Docling failed")
    return _empty_document(file_path, f"unsupported_{ext}", warnings)


def _empty_document(file_path: str, reason: str, warnings: list[str] | None = None) -> ResumeDocument:
    logger.warning("[Parser] Returning empty document for '%s' — reason: %s", file_path, reason)
    return ResumeDocument(
        source_path=file_path,
        parser_used=reason,
        parser_warnings=warnings or [f"All parsers failed: {reason}"],
    )
